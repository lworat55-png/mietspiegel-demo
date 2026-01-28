import streamlit as st # Das Werkzeug für die Webseite
from openai import OpenAI # Die Verbindung zu OpenAI
from docx import Document # Werkzeug zum Erstellen von Word-Dokumenten
import io # Hilfsmittel für den Dateiversand

# 1. SETUP: Wir holen uns die geheimen Daten (Prompt & API-Key) aus dem Tresor
# Diese Daten werden NICHT im Code stehen, sondern später in den Web-Einstellungen hinterlegt.
api_key = st.secrets["OPENAI_API_KEY"]
geheimer_system_prompt = st.secrets["SYSTEM_PROMPT"]

client = OpenAI(api_key=api_key)

# 2. DAS INTERFACE: Was der Anwalt auf der Webseite sieht
st.set_page_config(page_title="KI-Mietanalyst Demo", layout="centered")
st.title("⚖️ KI-Mietwert-Analyst")
st.info("Demo-Version für Anwaltskanzleien: Fallprüfung & Schriftsatz-Entwurf")

# Eingabefeld für den Sachverhalt
user_input = st.text_area("Geben Sie hier die Falldaten ein (z.B. Eckdaten Mietvertrag & Mietspiegel-Werte):", height=200)

# 3. DIE LOGIK: Was passiert beim Klick auf den Button?
if st.button("Analyse & Dokument generieren"):
    if not user_input:
        st.warning("Bitte geben Sie zuerst Daten ein.")
    else:
        with st.spinner("KI analysiert den Fall und erstellt das Dokument..."):
            # Hier schicken wir die Daten an OpenAI, inkl. deinem geheimen Prompt
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": geheimer_system_prompt},
                    {"role": "user", "content": user_input}
                ]
            )
            
            ergebnis_text = response.choices[0].message.content
            
            # Ergebnis auf der Webseite anzeigen
            st.subheader("Analyse-Ergebnis")
            st.markdown(ergebnis_text)
            
            # 4. DOKUMENT-ERSTELLUNG: Wir packen den Text in eine Word-Datei
            doc = Document()
            doc.add_heading('Analyse & Schriftsatz-Entwurf', 0)
            doc.add_paragraph(ergebnis_text)
            
            # Datei im Speicher "auffangen" für den Download
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.download_button(
                label="📄 Dokument als Word-Datei herunterladen",
                data=buffer,
                file_name="Mietwert_Analyse.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
